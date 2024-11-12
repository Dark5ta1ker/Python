import datetime
import locale

from docx import Document

locale.setlocale(locale.LC_ALL, 'ru_RU') #Устанавливаем локаль на русский

class DocProcessing:
    
    def __init__(self, doc_path):
        self.doc = Document(doc_path)
        self.date_format = "%d.%m.%Y"
        self.current_table_index = 0  # Индекс текущей таблицы
        self.current_cell_index = 0   # Индекс текущей ячейки в таблице

    def table_fill(self, user_input):
        """Заполняет таблицу в зависимости от номера таблицы и данных пользователя"""
        table = self.doc.tables[self.current_table_index]
        
        match self.current_table_index:
            case 0:                         #Дата
                converted_date = self.convert_date_format(user_input)
                self.clear(table.rows[0].cells[1]).text = '«' + converted_date[0] + '»'
                self.clear(table.rows[0].cells[2]).text = converted_date[1]
                self.clear(table.rows[0].cells[3]).text = converted_date[2] + 'г.'
            
            case 1:                         #№ Акта
                self.clear(table.rows[0].cells[1]).text = user_input
            
            case 2:                         #Изделие, зав№, № заказа
                cells = [
                    table.rows[0].cells[1], 
                    table.rows[0].cells[3], 
                    table.rows[1].cells[3]
                ]                
                self.format_text(self.clear(cells[self.current_cell_index]).paragraphs[0].add_run(user_input))
            
            case 3:                         #Адрес неисправности
                part1, part2, part3, part4 = self.split_string_by_spaces(user_input, 40, 4, 271) 
            
                self.format_text(self.clear(table.rows[0].cells[1]).paragraphs[0].add_run(part1))
                self.format_text(self.clear(table.rows[2].cells[0]).paragraphs[0].add_run(part2))
                self.format_text(self.clear(table.rows[4].cells[0]).paragraphs[0].add_run(part3))
                self.format_text(self.clear(table.rows[5].cells[0]).paragraphs[0].add_run(part4))
                return "Ведите адрес неисправности"
            
            case 4:                         #Признаки неисправности изделия
                part1, part2, part3, part4 = self.split_string_by_spaces(user_input, 35, 4, 266)
            
                self.format_text(self.clear(table.rows[0].cells[1]).paragraphs[0].add_run(part1))
                self.format_text(self.clear(table.rows[2].cells[0]).paragraphs[0].add_run(part2))
                self.format_text(self.clear(table.rows[4].cells[0]).paragraphs[0].add_run(part3))
                self.format_text(self.clear(table.rows[5].cells[0]).paragraphs[0].add_run(part4))
                
            case 5:                         #Предполагаемая причина неисправности
                part1, part2, part3, part4 = self.split_string_by_spaces(user_input, 31, 3, 185)
            
                self.format_text(self.clear(table.rows[1].cells[1]).paragraphs[0].add_run(part1))
                self.format_text(self.clear(table.rows[2].cells[0]).paragraphs[0].add_run(part2))
                self.format_text(self.clear(table.rows[3].cells[0]).paragraphs[0].add_run(part3))
            
            case 6:                         #Другие неисправности в изделии
                part1, part2, part3, part4 = self.split_string_by_spaces(user_input, 77, 2, 154)
            
                self.format_text(self.clear(table.rows[2].cells[0]).paragraphs[0].add_run(part1))
                self.format_text(self.clear(table.rows[3].cells[0]).paragraphs[0].add_run(part2))
                
            case 7:                         #Сколько отработало
                cells = [
                    table.rows[1].cells[1],
                    table.rows[1].cells[3],
                    table.rows[2].cells[1],
                    table.rows[2].cells[3]
                ]
                self.format_text(self.clear(cells[self.current_cell_index]).paragraphs[0].add_run(user_input))

            case 8:                         #Как сломалось
                part1, part2, part3, part4 = self.split_string_by_spaces(user_input, 16, 3, 170)
                
                self.format_text(self.clear(table.rows[1].cells[1]).paragraphs[0].add_run(part1))
                self.format_text(self.clear(table.rows[2].cells[0]).paragraphs[0].add_run(part2))
                self.format_text(self.clear(table.rows[3].cells[0]).paragraphs[0].add_run(part3))
                
            case 9:                         #Внешние факторы
                cells = [
                    table.rows[0].cells[3],
                    table.rows[2].cells[1],
                    table.rows[4].cells[2]
                ]
                self.format_text(self.clear(cells[self.current_cell_index]).paragraphs[0].add_run(user_input))
            
            case 10:                        #Признаки неисправности отказавшего элемента
                part1, part2, part3, part4 = self.split_string_by_spaces(user_input, 77, 2, 154)
                
                self.format_text(self.clear(table.rows[1].cells[1]).paragraphs[0].add_run(part1))
                self.format_text(self.clear(table.rows[3].cells[0]).paragraphs[0].add_run(part2))
            
            case 11:                        #Обстоятельства возникновения неисправности
                self.format_text(self.clear(table.rows[2].cells[0]).paragraphs[0].add_run(self.valid_string_length(user_input, 77)))
            
            case 12:                        #Какими мерами восстановлена работоспособность
                part1, part2, part3, part4 = self.split_string_by_spaces(user_input, 13, 3, 167)  
                
                self.format_text(self.clear(table.rows[1].cells[1]).paragraphs[0].add_run(part1))
                self.format_text(self.clear(table.rows[2].cells[0]).paragraphs[0].add_run(part2)) 
                self.format_text(self.clear(table.rows[4].cells[0]).paragraphs[0].add_run(part3))
            
            case 13:                        #Сколько искали/чинили
                cells = [
                    table.rows[1].cells[1],
                    table.rows[1].cells[3],
                    table.rows[2].cells[1],
                    table.rows[2].cells[3]
                ]
                self.format_text(self.clear(cells[self.current_cell_index]).paragraphs[0].add_run(user_input))
            
            case 14:                        #Этап работы
                part1, part2, part3, part4 = self.split_string_by_spaces(user_input, 50, 2, 127)
                
                self.format_text(self.clear(table.rows[0].cells[1]).paragraphs[0].add_run(part1))
                self.format_text(self.clear(table.rows[2].cells[0]).paragraphs[0].add_run(part2)) 
            
            case 15:                        #Экземпляров акта составлено
                part1, part2, part3, part4 = self.split_string_by_spaces(user_input, 3, 2, 80)
                
                self.format_text(self.clear(table.rows[0].cells[1]).paragraphs[0].add_run(part1))
                self.format_text(self.clear(table.rows[1].cells[0]).paragraphs[0].add_run(part2))
            
            case 16:                        #Отказавшие элементы
                part1, part2, part3, part4 = self.split_string_by_spaces(user_input, 16, 2, 122)
                
                self.format_text(self.clear(table.rows[0].cells[1]).paragraphs[0].add_run(part1))
                self.format_text(self.clear(table.rows[2].cells[0]).paragraphs[0].add_run(part2))
             
            case 17:                        #Направлены в
                self.format_text(self.clear(table.rows[0].cells[1]).paragraphs[0].add_run(user_input))
            
            case 18:                        #Подпись, дата
                cells = [
                    table.rows[2].cells[2],
                    table.rows[2].cells[4]
                ]
                self.format_text(self.clear(cells[self.current_cell_index]).paragraphs[0].add_run(user_input))
        return "Нет подсказки для этой таблицы"

    def get_current_table_cells(self):                                        #Возвращает список ячеек для текущей таблицы
        match self.current_table_index:
            case 2:
                return [self.doc.tables[2].rows[0].cells[1], self.doc.tables[2].rows[0].cells[3], self.doc.tables[2].rows[1].cells[3]]

            case 7:
                return [self.doc.tables[7].rows[1].cells[1], self.doc.tables[7].rows[1].cells[3], self.doc.tables[7].rows[2].cells[1], self.doc.tables[7].rows[2].cells[3]]
 
            case 9:
                return [self.doc.tables[9].rows[0].cells[3], self.doc.tables[9].rows[2].cells[1], self.doc.tables[9].rows[4].cells[2]]
 
            case 13:
                return [self.doc.tables[13].rows[1].cells[1], self.doc.tables[13].rows[1].cells[3], self.doc.tables[13].rows[2].cells[1], self.doc.tables[13].rows[2].cells[3]]

            case 18:
                return [self.doc.tables[18].rows[2].cells[1], self.doc.tables[18].rows[2].cells[2]]

        return []
    
    def move_forward(self):                                                    #Перемещает на одну ячейку вперёд, или переходит к следующей таблице.
        cells_in_table = len(self.get_current_table_cells())
        if self.current_cell_index < cells_in_table - 1:
            # Переход к следующей ячейке в таблице
            self.current_cell_index += 1
        else:
            # Переход к следующей таблице и сброс индекса ячейки
            if self.current_table_index < len(self.doc.tables) - 1:
                self.current_table_index += 1
                self.current_cell_index = 0  # Начало новой таблицы
            else:
                print("Вы достигли конца документа.")

    def move_backward(self):                                                   #Перемещает на одну ячейку назад, или возвращает к предыдущей таблице.
        if self.current_cell_index > 0:
            # Переход к предыдущей ячейке
            self.current_cell_index -= 1
        else:
            # Переход к предыдущей таблице
            if self.current_table_index > 0:
                self.current_table_index -= 1
                # Переход к последней ячейке предыдущей таблицы
                self.current_cell_index = len(self.get_current_table_cells()) - 1
            else:
                print("Вы достигли начала документа.")
        if self.current_cell_index > 0:
            self.current_cell_index -= 1

    def reset_cell_index(self):                                                #Сброс индекса ячейки при смене таблицы
        self.current_cell_index = 0

    def get_hint(self):
        """Возвращает подсказку для текущей таблицы и ячейки без изменения данных"""
        # Определим подсказки для каждой таблицы по её индексу
        match self.current_table_index:
            case 0:
                return "Введите дату в формате дд.мм.гггг"
            case 1:
                return "Введите номер акта"
            case 2:
                hints = [
                    "Введите название изделия", 
                    "Введите заводской номер изделия", 
                    "Введите номер заказа изделия"
                ]
                return hints[self.current_cell_index]
            case 3:
                return "Введите адрес неисправности"
            case 4:
                return "Признаки неисправности изделия"
            case 5:
                return "Предполагаемая причина неисправности:"
            case 6:
                return "Другие неисправности в изделии:"
            case 7:
                hints = [
                    "Введите, сколько часов отработало изделие (прибор)",
                    "Введите, сколько минут отработало изделие (прибор)",
                    "Введите, сколько часов отработал отказавший элемент",
                    "Введите, сколько минут отработал отказавший элемент"
                ]
                return hints[self.current_cell_index]
            case 8:
                return "Режим работы блока/прибора перед моментом возникновения неисправности:"
            case 9:
                hints = [
                    "Введите напряжение источника электропитания",
                    "Введите климатические условия",
                    "Введите механические воздействия"
                ]
                return hints[self.current_cell_index]
            case 10:
                return "Признаки неисправности отказавшего элемента"                
            case 11:
                return "Обстоятельства возникновения неисправности:"
            case 12:
                return "Какими мерами восстановлена работоспособность"
            case 13:
                hints = [
                    "Введите, сколько часов искали неисправность",
                    "Введите, сколько минут искали неисправность",
                    "Введите, сколько часов чинили неисправность",
                    "Введите, сколько минут чинили неисправность"
                ]
                return hints[self.current_cell_index]
            case 14:
                return "Этап работы"
            case 15:
                return "Экземпляров акта составлено и направлено в:"
            case 16:
                return "Отказавшие элементы"
            case 17:
                return "Направлены в:"
            case 18:
                hints = [
                    "Введите дату составления акта",
                    "Введите фамилию и инициалы составителя акта"
                ]
                return hints[self.current_cell_index]   

        return "Нет подсказки для этой таблицы"

    def split_string_by_spaces(self, text, cell1_length, strings, max_length): #Перенос строк при переполнении ячейки. На 37 ебаных строк кода ушло 5 часов 20 минут
        L = 77  # Максимальная длина для каждого сегмента
        if len(text) > max_length:
            print('Текст слишком длинный, укоротите его на', len(text) - max_length, 'символов')
            return "", "", "", ""

        # Поиск индексов для разделения текста
        space_index1 = text.rfind(' ', 0, cell1_length)
        if space_index1 == -1 or space_index1 == cell1_length - 1:
            space_index1 = cell1_length

        space_index2 = text.rfind(' ', space_index1 + 1, space_index1 + L)
        if space_index2 == -1 or (len(text) - space_index1 <= L):
            space_index2 = space_index1 + L

        space_index3 = text.rfind(' ', space_index2 + 1, space_index2 + L)
        if space_index3 == -1 or (len(text) - space_index2 <= L):
            space_index3 = space_index2 + L

        # Разделение текста на части без смещения, чтобы не терять буквы
        part1 = text[:space_index1]
        part2 = text[space_index1:].strip()[:space_index2 - space_index1].strip()  # Используем полный диапазон
        part3 = text[space_index2:].strip()[:space_index3 - space_index2].strip()
        part4 = text[space_index3:].strip()[:L].strip() if len(text) > space_index3 else ""

        # Возвращение в зависимости от количества строк
        if strings == 2:
            return part1, part2, "", ""
        elif strings == 3:
            return part1, part2, part3, ""
        elif strings == 4:
            return part1, part2, part3, part4
        else:
            return part1, part2, part3, part4
  
    def convert_date_format(self, date_str):                                   #Конвертируем дату формата дд.мм.гггг в дд.месяц.гггг
        day, month, year = date_str.split('.')
        date_obj = datetime.datetime(int(year), int(month), int(day))
        month_name = date_obj.strftime('%B')
        return day, month_name, year        

    def valid_string_length(self, user_input, max_length):                     #Проверяем допустимую длину строки
        while True:
            if len(user_input) <= max_length:
                return user_input
            else:
                print('Слишком длинная строка. Введите строку короче на ' + str(max_length - len((user_input))) + 'символов, или не вводите ничего и напишите от руки')

    def format_text(self, run):                                                #Форматируем текст. Все гениальное просто...
        run.font.name = 'GOST type B'
        run.font.italic = True
        
    def clear(self, cell):                                                     #Очистка ячейки перед записью
        cell.text = ''
        return cell

    def save_document(self, new_path="new.docx"):
        """Сохраняет изменения в документ."""
        self.doc.save(new_path)
        print(f"Документ сохранен как {new_path}")




